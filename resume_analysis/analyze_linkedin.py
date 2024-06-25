from bs4 import BeautifulSoup
import json
from docx import Document

def parse_linkedin_profile(html_content):
    """
    Parse LinkedIn profile HTML content and extract relevant data.

    Args:
        html_content (str): HTML content of the LinkedIn profile.

    Returns:
        dict: Structured profile data.
    """
    soup = BeautifulSoup(html_content, 'html.parser')

    profile_data = {
        'name': '',
        'headline': '',
        'summary': '',
        'experience': [],
        'education': [],
        'skills': [],
        'certifications': []
    }

    # Extract name
    name_tag = soup.find('h1', {'class': 'top-card-layout__title'})
    if name_tag:
        profile_data['name'] = name_tag.get_text(strip=True)

    # Extract headline
    headline_tag = soup.find('h2', {'class': 'top-card-layout__headline'})
    if headline_tag:
        profile_data['headline'] = headline_tag.get_text(strip=True)

    # Extract summary
    summary_tag = soup.find('section', {'id': 'summary'})
    if summary_tag:
        profile_data['summary'] = summary_tag.get_text(strip=True)

    # Extract experience
    experience_section = soup.find('section', {'id': 'experience-section'})
    if experience_section:
        for exp in experience_section.find_all('li', {'class': 'result-card'}):
            title = exp.find('h3', {'class': 'result-card__title'}).get_text(strip=True)
            company = exp.find('h4', {'class': 'result-card__subtitle'}).get_text(strip=True)
            date_range = exp.find('span', {'class': 'result-card__date-range'}).get_text(strip=True)
            location = exp.find('span', {'class': 'result-card__location'}).get_text(strip=True)
            profile_data['experience'].append({
                'title': title,
                'company': company,
                'date_range': date_range,
                'location': location
            })

    # Extract education
    education_section = soup.find('section', {'id': 'education-section'})
    if education_section:
        for edu in education_section.find_all('li', {'class': 'result-card'}):
            school = edu.find('h3', {'class': 'result-card__title'}).get_text(strip=True)
            degree = edu.find('span', {'class': 'result-card__degree'}).get_text(strip=True)
            field_of_study = edu.find('span', {'class': 'result-card__field-of-study'}).get_text(strip=True)
            date_range = edu.find('span', {'class': 'result-card__date-range'}).get_text(strip=True)
            profile_data['education'].append({
                'school': school,
                'degree': degree,
                'field_of_study': field_of_study,
                'date_range': date_range
            })

    # Extract skills
    skills_section = soup.find('section', {'id': 'skills-section'})
    if skills_section:
        for skill in skills_section.find_all('li', {'class': 'result-card'}):
            skill_name = skill.find('span', {'class': 'result-card__skill-name'}).get_text(strip=True)
            profile_data['skills'].append(skill_name)

    # Extract certifications
    certifications_section = soup.find('section', {'id': 'certifications-section'})
    if certifications_section:
        for cert in certifications_section.find_all('li', {'class': 'result-card'}):
            cert_name = cert.find('h3', {'class': 'result-card__title'}).get_text(strip=True)
            issuing_organization = cert.find('span', {'class': 'result-card__issuing-organization'}).get_text(strip=True)
            date_range = cert.find('span', {'class': 'result-card__date-range'}).get_text(strip=True)
            profile_data['certifications'].append({
                'cert_name': cert_name,
                'issuing_organization': issuing_organization,
                'date_range': date_range
            })

    return profile_data

def format_resume(profile_data, output_file):
    """
    Format the extracted profile data into a professional resume layout and save it as a DOCX file.

    Args:
        profile_data (dict): Structured profile data.
        output_file (str): Path to the output DOCX file.
    """
    document = Document()

    # Add name and headline
    document.add_heading(profile_data['name'], level=1)
    document.add_paragraph(profile_data['headline'])

    # Add summary
    document.add_heading('Summary', level=2)
    document.add_paragraph(profile_data['summary'])

    # Add experience
    document.add_heading('Experience', level=2)
    for exp in profile_data['experience']:
        document.add_heading(exp['title'], level=3)
        document.add_paragraph(f"{exp['company']} - {exp['date_range']}")
        document.add_paragraph(exp['location'])

    # Add education
    document.add_heading('Education', level=2)
    for edu in profile_data['education']:
        document.add_heading(edu['school'], level=3)
        document.add_paragraph(f"{edu['degree']} in {edu['field_of_study']} - {edu['date_range']}")

    # Add skills
    document.add_heading('Skills', level=2)
    document.add_paragraph(', '.join(profile_data['skills']))

    # Add certifications
    document.add_heading('Certifications', level=2)
    for cert in profile_data['certifications']:
        document.add_heading(cert['cert_name'], level=3)
        document.add_paragraph(f"Issued by {cert['issuing_organization']} - {cert['date_range']}")

    # Save the document
    document.save(output_file)

# Example usage
if __name__ == "__main__":
    with open('linkedin_profile.html', 'r') as file:
        html_content = file.read()

    profile_data = parse_linkedin_profile(html_content)
    print(json.dumps(profile_data, indent=4))

    format_resume(profile_data, 'resume.docx')
